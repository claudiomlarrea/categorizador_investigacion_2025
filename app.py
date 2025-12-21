import streamlit as st
import re, json, io
import pandas as pd
import unicodedata
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF opcional
try:
    import pdfplumber
    HAVE_PDF = True
except Exception:
    HAVE_PDF = False

st.set_page_config(page_title="Valorador de CV - UCCuyo (DOCX/PDF)", layout="wide")
st.title("Universidad Católica de Cuyo — Valorador de CV Docente")
st.caption("Incluye exportación a Excel y Word + categoría automática según puntaje total.")

@st.cache_data(show_spinner=False)
def load_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        st.error(f"criteria.json inválido: {e.msg} (línea {e.lineno}, columna {e.colno}).")
        st.info("Tip: revisá comillas, comas finales y backslashes en regex (en JSON deben ser \\\\).")
        st.stop()
    except FileNotFoundError:
        st.error("No se encontró criteria.json en el repositorio (debe estar en la misma carpeta que app.py).")
        st.stop()
    except Exception as e:
        st.error(f"Error leyendo criteria.json: {e}")
        st.stop()

criteria = load_json("criteria.json")

# Debug SOLO en el cuerpo (NO sidebar)
DEBUG = st.checkbox("Debug (mostrar texto y entradas detectadas)", value=False)

# =========================
# Extracción de texto
# =========================
def extract_text_docx(file):
    doc = DocxDocument(file)
    text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            text += "\n" + " | ".join(c.text for c in row.cells)
    return text

def extract_text_pdf(file):
    if not HAVE_PDF:
        raise RuntimeError("Falta pdfplumber. Agregalo en requirements.txt: pdfplumber")
    chunks = []
    with pdfplumber.open(file) as pdf:
        for p in pdf.pages:
            chunks.append(p.extract_text() or "")
    return "\n".join(chunks)

# =========================
# Helpers
# =========================
def _strip_accents(s: str) -> str:
    if not s:
        return ""
    s = unicodedata.normalize("NFKD", s)
    return "".join(ch for ch in s if not unicodedata.combining(ch))

def match_count(pattern, text):
    """
    Regex global (para secciones no parseadas) con fallback por normalización
    (evita fallos por Año/Ano, finalización/finalizacion, etc.)
    """
    if not pattern:
        return 0
    try:
        m1 = re.findall(pattern, text, flags=re.IGNORECASE | re.UNICODE)
        if m1:
            return len(m1)
    except re.error:
        return 0

    # Fallback: quitar acentos tanto a patrón como a texto
    try:
        text2 = _strip_accents(text)
        pat2 = _strip_accents(pattern)
        return len(re.findall(pat2, text2, flags=re.IGNORECASE | re.UNICODE))
    except re.error:
        return 0

def clip(v, cap):
    return min(v, cap) if cap else v

def normalize_spaces(s: str) -> str:
    s = s.replace("\u00A0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def obtener_categoria(total, criteria_dict):
    categorias = criteria_dict.get("categorias", {})
    mejor_clave = "Sin categoría"
    mejor_desc = ""
    mejor_min = None
    for clave, info in categorias.items():
        min_pts = info.get("min_points", 0)
        if total >= min_pts and (mejor_min is None or min_pts > mejor_min):
            mejor_min = min_pts
            mejor_clave = clave
            mejor_desc = info.get("descripcion", "")
    return mejor_clave, mejor_desc

# ==========================================================
# 1) Recorte robusto de la sección "FORMACION ACADÉMICA"
# ==========================================================
FORMACION_HEADERS = [
    r"FORMACI[ÓO]N ACAD[ÉE]MICA",
    r"FORMACION ACADEMICA",
    r"FORMACI[ÓO]N\s+ACAD[ÉE]MICA",
    r"FORMACI[ÓO]N\s+ACAD[ÉE]MICA\s+Y\s+COMPLEMENTARIA",
    r"FORMACION\s+ACADEMICA\s+Y\s+COMPLEMENTARIA",
]

NEXT_SECTION_MARKERS = [
    r"\n\s*FORMACI[ÓO]N\s+DE\s+RECURSOS\s+HUMANOS\b",
    r"\n\s*RECURSOS\s+HUMANOS\b",
    r"\n\s*RRHH\b",
    r"\n\s*ANTECEDENTES\b",
    r"\n\s*PRODUCCI[ÓO]N\b",
    r"\n\s*PUBLICACIONES\b",
    r"\n\s*ACTIVIDADES\b",
    r"\n\s*EXPERIENCIA\b",
    r"\n\s*CARGOS\b",
]

def extract_formacion_academica_block(full_text: str) -> str:
    txt = normalize_spaces(full_text)
    start_idx = None
    for h in FORMACION_HEADERS:
        m = re.search(h, txt, flags=re.IGNORECASE)
        if m:
            start_idx = m.end()
            break
    if start_idx is None:
        return ""

    tail = txt[start_idx:]
    end_idx = len(tail)
    for mk in NEXT_SECTION_MARKERS:
        m2 = re.search(mk, tail, flags=re.IGNORECASE)
        if m2:
            end_idx = min(end_idx, m2.start())
    return tail[:end_idx].strip()

# ==========================================================
# 2) Parseo por entradas + FINALIZACIÓN EXPLÍCITA
#    (NO usar rangos de años como evidencia de finalización)
# ==========================================================
RE_IN_PROGRESS = re.compile(
    r"\b(Actualidad|En\s+curso|Cursando|Actualmente|Vigente|En\s+desarrollo|Hasta\s+la\s+actualidad|A\s+la\s+fecha)\b",
    re.IGNORECASE
)

RE_FINISH_YEAR = re.compile(
    r"A[nñ]o\s+de\s+(finalizaci[oó]n|obtenci[oó]n|graduaci[oó]n)\s*:\s*([0-3]?\d/\d{4}|\d{4})",
    re.IGNORECASE
)

RE_SITUACION_COMPLETO = re.compile(
    r"Situaci[oó]n\s+del\s+nivel\s*:\s*Completo",
    re.IGNORECASE
)

# Marcadores explícitos (tolerantes) de finalización/egreso
RE_COMPLETION_CUES = re.compile(
    r"\b(finalizad[oa]|egresad[oa]|graduad[oa]|t[ií]tulo\s+obtenido|t[ií]tulo\s+otorgado|complet(?:o|ada))\b",
    re.IGNORECASE
)

RE_BECARIO_CONTEXT = re.compile(
    r"\b(becari[oa]s?|beca|direcci[oó]n|co[- ]?direcci[oó]n|tesista|investigador/a|investigador)\b",
    re.IGNORECASE
)

# ✅ FIX CLAVE: incluir "Licenciados" (plural) y "Profesor Universitario"
RE_ENTRY_START = re.compile(
    r"^\s*(?:[-•·*]\s*)?"
    r"(Doctorado|Doctor\s+en|Doctor\s+de\s+la\s+Universidad|Maestr[ií]a|Mag[ií]ster|"
    r"Especializaci[oó]n|Especialista|"
    r"Profesorado|Profesor\s+Universitario|Profesor\s+en|"
    r"Licenciatura|Licenciad[oa]s?|T[eé]cnica\s+Universitaria|Tecnicatura|"
    r"Contador|Contadora|Contadur[ií]a|"
    r"Abogado|Abogada|Ingenier|Bioqu[ií]mic|M[eé]dic|Farmac[eé]utic|Arquitect|Odont[oó]log)\b",
    re.IGNORECASE
)

# POSDOC: también requiere evidencia explícita (no basta “Actualidad”)
RE_POSDOC_ENTRY = re.compile(r"(?ims)^(Posdoctorado|Postdoctorado)\b[\s\S]{0,1600}", re.IGNORECASE)

def split_entries(block: str) -> list[str]:
    if not block:
        return []
    lines = [l.strip() for l in block.split("\n")]
    lines = [l for l in lines if l and l.lower() != "null"]

    entries = []
    buf = []
    for line in lines:
        if RE_ENTRY_START.search(line) and buf:
            entries.append("\n".join(buf).strip())
            buf = [line]
        else:
            buf.append(line)

    if buf:
        entries.append("\n".join(buf).strip())

    # fallback SOLO por títulos principales (sin posdoc)
    if len(entries) == 1 and len(entries[0]) > 1500:
        parts = re.split(
            r"(?i)(?=Doctorado\b|Maestr[ií]a\b|Especializaci[oó]n\b|Licenciatura\b|Licenciad[oa]s?\b|"
            r"T[eé]cnica\s+Universitaria\b|Tecnicatura\b|Profesorado\b|Profesor\s+Universitario\b)",
            entries[0]
        )
        entries = [p.strip() for p in parts if p.strip()]

    return entries

def entry_is_completed(entry: str) -> bool:
    # Si dice en curso/actualidad => NO
    if RE_IN_PROGRESS.search(entry):
        return False
    # Evidencia fuerte => SI
    if RE_FINISH_YEAR.search(entry):
        return True
    if RE_SITUACION_COMPLETO.search(entry):
        return True
    if RE_COMPLETION_CUES.search(entry):
        return True
    # Si solo hay años sueltos o rangos => NO (regla dura anti-falsos positivos)
    return False

def get_finish_token(entry: str) -> str:
    m = RE_FINISH_YEAR.search(entry)
    if m:
        return m.group(2).strip()
    if RE_SITUACION_COMPLETO.search(entry):
        return "COMPLETO"
    if RE_COMPLETION_CUES.search(entry):
        return "FINALIZADO"
    return ""

def get_first_line_title(entry: str) -> str:
    lines = [l.strip() for l in entry.split("\n") if l.strip()]
    for l in lines:
        if l.lower() == "null":
            continue
        return l
    return (lines[0] if lines else "").strip()

def get_institution_hint(entry: str) -> str:
    lines = [l.strip() for l in entry.split("\n") if l.strip()]
    for l in lines[:10]:
        if re.search(r"\b(UNIVERSIDAD|FACULTAD|INSTITUTO|SEDE|CONICET)\b", l, re.IGNORECASE):
            return l
    return ""

def norm_key(s: str) -> str:
    s = (s or "").lower().strip()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"[\"'`´]", "", s)
    return s

def classify_entry(entry: str) -> str:
    if re.search(r"\bDoctorado\b|\bDoctor\s+en\b|\bDoctor\s+de\s+la\s+Universidad\b", entry, re.IGNORECASE):
        return "doctorado"
    if re.search(r"\bMaestr[ií]a\b|\bMag[ií]ster\b", entry, re.IGNORECASE):
        return "maestria"
    if re.search(r"\bEspecializaci[oó]n\b|\bEspecialista\b", entry, re.IGNORECASE):
        return "especializacion"

    if re.search(r"\bPos\s*graduad[oa]\b|\bPos\s*grado\b|\bPosgrado\b", entry, re.IGNORECASE):
        return "otro"

    # ✅ FIX CLAVE: Profesor Universitario cuenta como profesorado
    if re.search(r"\bProfesorado\b|\bProfesor\s+en\b|\bProfesor\s+Universitario\b", entry, re.IGNORECASE):
        return "profesorado"

    # ✅ FIX CLAVE: incluir Licenciados (plural)
    if re.search(
        r"\b(Licenciatura|Licenciad[oa]s?|T[eé]cnica\s+Universitaria|Tecnicatura|"
        r"Contador|Contadora|Contadur[ií]a|Abogado|Abogada|Ingenier|Bioqu[ií]mic|M[eé]dic|Farmac[eé]utic|Arquitect|Odont[oó]log)\b",
        entry,
        re.IGNORECASE
    ):
        return "grado"

    return "otro"

def count_posdoc_explicit(block: str) -> int:
    """
    Posdoc SOLO si:
    - existe una entrada que comienza con Posdoctorado/Postdoctorado
    - y tiene evidencia explícita de finalización (misma regla dura)
    """
    if not block:
        return 0
    matches = []
    for m in RE_POSDOC_ENTRY.finditer(block):
        chunk = m.group(0)
        # Excluir si parece RRHH/becas
        if RE_BECARIO_CONTEXT.search(chunk):
            continue
        # Debe ser "entrada" real (primera línea ya es posdoc), y además finalización explícita
        if entry_is_completed(chunk):
            matches.append(chunk)

    seen = set()
    for x in matches:
        k = norm_key(re.sub(r"\s+", " ", x)[:400])
        seen.add(k)
    return len(seen)

def counts_from_formacion(block: str) -> dict:
    entries = split_entries(block)
    seen = set()

    counts = {
        "doctorado": 0,
        "maestria": 0,
        "especializacion": 0,
        "grado": 0,
        "profesorado": 0,
        "posdoc": 0,
    }

    counts["posdoc"] = count_posdoc_explicit(block)

    for e in entries:
        tipo = classify_entry(e)
        if tipo not in counts or tipo == "posdoc":
            continue

        if not entry_is_completed(e):
            continue

        titulo = get_first_line_title(e)
        fin = get_finish_token(e)
        inst = get_institution_hint(e)

        key = (tipo, norm_key(titulo), norm_key(inst), norm_key(fin))
        if key in seen:
            continue

        seen.add(key)
        counts[tipo] += 1

    return counts

# =========================
# UI
# =========================
uploaded = st.file_uploader("Cargar CV (.docx o .pdf)", type=["docx", "pdf"])

if uploaded:
    ext = uploaded.name.split(".")[-1].lower()
    try:
        raw_text = extract_text_docx(uploaded) if ext == "docx" else extract_text_pdf(uploaded)
    except Exception as e:
        st.error(str(e))
        st.stop()

    raw_text = normalize_spaces(raw_text)
    st.success(f"Archivo cargado: {uploaded.name}")

    form_block = extract_formacion_academica_block(raw_text)
    form_counts = counts_from_formacion(form_block)

    if DEBUG:
        with st.expander("Ver texto extraído (debug)"):
            st.text_area("Texto", raw_text, height=240)

        with st.expander("Ver sección de Formación académica (debug)"):
            st.text_area("FORMACIÓN ACADÉMICA (recorte)", form_block if form_block else "[No se encontró la sección]", height=240)

        with st.expander("Ver entradas detectadas en Formación (debug avanzado)"):
            entries_dbg = split_entries(form_block)
            st.write(f"Entradas detectadas: {len(entries_dbg)}")
            for i, ent in enumerate(entries_dbg[:60], start=1):
                st.markdown(f"**Entrada {i}** — tipo: `{classify_entry(ent)}` — finalizado: `{entry_is_completed(ent)}`")
                st.code(ent[:1200])
            st.write(f"Posdoc explícitos finalizados detectados: {count_posdoc_explicit(form_block)}")
            st.write("Conteos formacion:", form_counts)

    results = {}
    total = 0.0

    # =========================
    # Cálculo de puntajes por sección
    # =========================
    for section, cfg in criteria["sections"].items():
        st.markdown(f"### {section}")
        rows = []
        subtotal_raw = 0.0

        for item, icfg in cfg.get("items", {}).items():
            pattern = icfg.get("pattern", "")
            c = None

            # Overrides SOLO para Formación académica y complementaria
            if section.lower().startswith("formación académica") or section.lower().startswith("formacion academica"):
                item_l = item.lower()

                if "doctorado" in item_l:
                    c = form_counts.get("doctorado", 0)
                elif "maestr" in item_l or "magíster" in item_l or "magister" in item_l:
                    c = form_counts.get("maestria", 0)
                elif "especializ" in item_l or "especialista" in item_l:
                    c = form_counts.get("especializacion", 0)
                elif "título de grado" in item_l or "titulo de grado" in item_l or item_l.strip() == "grado":
                    c = form_counts.get("grado", 0)
                elif "profesorado" in item_l or "docencia universitaria" in item_l:
                    c = form_counts.get("profesorado", 0)
                elif re.search(r"\bposdoc\b|\bpostdoc\b|\bposdoctorad\b|\bpostdoctorad\b", item_l):
                    c = form_counts.get("posdoc", 0)

            # 🔒 Bloqueo extra: evitar conteo de títulos fuera de Formación
            if c is None:
                item_l = item.lower()
                es_titulo = bool(re.search(
                    r"\b(doctorad|maestr|magister|especializ|posdoc|postdoc|posdoctor|postdoctor|t[ií]tulo de grado|grado|profesorado)\b",
                    item_l
                ))
                if es_titulo and not (section.lower().startswith("formación académica") or section.lower().startswith("formacion academica")):
                    c = 0

            # si no aplicó override/bloqueo -> regex global (con fallback por acentos)
            if c is None:
                c = match_count(pattern, raw_text)

            pts = clip(c * icfg.get("unit_points", 0), icfg.get("max_points", 0))
            rows.append({
                "Ítem": item,
                "Ocurrencias": c,
                "Puntaje (tope ítem)": pts,
                "Tope ítem": icfg.get("max_points", 0)
            })
            subtotal_raw += pts

        df = pd.DataFrame(rows)
        subtotal = clip(subtotal_raw, cfg.get("max_points", 0))
        st.dataframe(df, use_container_width=True)
        st.info(f"Subtotal {section}: {subtotal} / máx {cfg.get('max_points', 0)}")
        results[section] = {"df": df, "subtotal": subtotal}
        total += subtotal

    # =========================
    # Categoría
    # =========================
    clave_cat, desc_cat = obtener_categoria(total, criteria)
    categoria_label = "Sin categoría" if clave_cat == "Sin categoría" else f"Categoría {clave_cat}"

    st.markdown("---")
    st.subheader("Puntaje total y categoría")
    st.metric("Total acumulado", f"{total:.1f}")
    st.metric("Categoría alcanzada", categoria_label)
    if desc_cat:
        st.info(f"Descripción de la categoría: {desc_cat}")

    # =========================
    # Exportaciones
    # =========================
    st.markdown("---")
    st.subheader("Exportar resultados")

    # Excel
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
        for sec, data in results.items():
            data["df"].to_excel(writer, sheet_name=sec[:31], index=False)

        resumen = pd.DataFrame({
            "Sección": list(results.keys()),
            "Subtotal": [results[s]["subtotal"] for s in results]
        })
        resumen.loc[len(resumen)] = ["TOTAL", resumen["Subtotal"].sum()]
        resumen.loc[len(resumen)] = ["CATEGORÍA", categoria_label]
        resumen.to_excel(writer, sheet_name="RESUMEN", index=False)

    st.download_button(
        "Descargar Excel",
        data=out.getvalue(),
        file_name="valoracion_cv.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )

    # Word
    def export_word(results_dict, total_pts, cat_label, cat_desc):
        doc = DocxDocument()
        p = doc.add_paragraph("Universidad Católica de Cuyo — Secretaría de Investigación")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Informe de valoración de CV").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        doc.add_paragraph(f"Puntaje total: {total_pts:.1f}")
        doc.add_paragraph(f"Categoría alcanzada: {cat_label}")
        if cat_desc:
            doc.add_paragraph(cat_desc)

        for sec, data in results_dict.items():
            doc.add_heading(sec, level=2)
            df_sec = data["df"]
            if df_sec.empty:
                doc.add_paragraph("Sin ítems detectados.")
            else:
                tbl = doc.add_table(rows=1, cols=len(df_sec.columns))
                hdr = tbl.rows[0].cells
                for i, ccol in enumerate(df_sec.columns):
                    hdr[i].text = str(ccol)
                for _, row in df_sec.iterrows():
                    cells = tbl.add_row().cells
                    for i, ccol in enumerate(df_sec.columns):
                        cells[i].text = str(row[ccol])
            doc.add_paragraph(f"Subtotal sección: {data['subtotal']:.1f}")

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    st.download_button(
        "Descargar informe Word",
        data=export_word(results, total, categoria_label, desc_cat),
        file_name="informe_valoracion_cv.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

else:
    st.info("Subí un archivo para iniciar la valoración.")
